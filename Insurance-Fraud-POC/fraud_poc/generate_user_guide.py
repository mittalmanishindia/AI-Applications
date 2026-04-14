from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def add_horizontal_line(paragraph):
    """Add a horizontal line below a paragraph"""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

def create_comprehensive_user_guide():
    document = Document()

    # --- Title Page ---
    title = document.add_heading('Life Insurance Fraud Detection System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = document.add_paragraph('Complete User Guide & Technical Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(100, 100, 100)

    version = document.add_paragraph('Version 2.0 - Gemini AI Enhanced')
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_run = version.runs[0]
    version_run.font.size = Pt(11)
    version_run.italic = True

    document.add_page_break()

    # --- Table of Contents ---
    document.add_heading('Table of Contents', level=1)
    toc_items = [
        "1. Introduction",
        "2. System Overview",
        "3. Installation & Setup",
        "4. User Interface Guide",
        "5. Agent Capabilities",
        "6. AI/LLM Integration",
        "7. Technical Architecture",
        "8. Sample Workflows",
        "9. Troubleshooting",
        "10. Future Enhancements"
    ]
    for item in toc_items:
        document.add_paragraph(item, style='List Bullet')

    document.add_page_break()

    # --- 1. Introduction ---
    document.add_heading('1. Introduction', level=1)
    intro_para = document.add_paragraph()
    intro_para.add_run("The Life Insurance Fraud Detection System").bold = True
    intro_para.add_run(
        " is an AI-powered agentic application designed to assist "
        "Special Investigation Units (SIU) in evaluating insurance claims for potential fraud. "
        "By analyzing policy details, claim context, historical patterns, and leveraging Google's "
        "Gemini 2.0 Flash-Lite AI model, the system provides real-time risk assessments with "
        "human-like reasoning capabilities."
    )

    document.add_heading('Purpose', level=2)
    document.add_paragraph(
        "This system serves as a decision-support tool for fraud investigators, helping to:"
    )
    document.add_paragraph('• Identify high-risk claims requiring manual review', style='List Bullet')
    document.add_paragraph('• Reduce false positives through contextual analysis', style='List Bullet')
    document.add_paragraph('• Provide explainable AI reasoning for audit trails', style='List Bullet')
    document.add_paragraph('• Accelerate claim processing for low-risk cases', style='List Bullet')

    # --- 2. System Overview ---
    document.add_heading('2. System Overview', level=1)
    
    document.add_heading('2.1 Architecture', level=2)
    document.add_paragraph(
        "The system operates on a lightweight, serverless architecture using Python and Streamlit. "
        "It does not require a traditional database, instead using structured JSON files to simulate "
        "enterprise data stores."
    )
    
    p = document.add_paragraph()
    p.add_run("Core Components:").bold = True
    document.add_paragraph('• User Interface (UI): A web-based dashboard built with Streamlit', style='List Bullet')
    document.add_paragraph('• AI Agent (Backend): Python logic orchestrating data fetching, rule evaluation, and LLM calls', style='List Bullet')
    document.add_paragraph('• Data Layer: JSON repositories for Claims, Policies, and Fraud Rules', style='List Bullet')
    document.add_paragraph('• LLM Integration: Google Gemini 2.0 Flash-Lite for natural language reasoning', style='List Bullet')

    document.add_heading('2.2 Data Model', level=2)
    document.add_paragraph("The system uses three primary data files:")
    
    table = document.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    headers = table.rows[0].cells
    headers[0].text = 'File'
    headers[1].text = 'Description'
    
    row1 = table.rows[1].cells
    row1[0].text = 'claims.json'
    row1[1].text = 'Contains claim details: amount, cause of death, location, documents'
    
    row2 = table.rows[2].cells
    row2[0].text = 'policies.json'
    row2[1].text = 'Policy holder profiles: coverage, start date, medical history flags'
    
    row3 = table.rows[3].cells
    row3[0].text = 'insurance_rules.json'
    row3[1].text = 'Configurable fraud detection rules with risk scores'

    # --- 3. Installation & Setup ---
    document.add_page_break()
    document.add_heading('3. Installation & Setup', level=1)
    
    document.add_heading('3.1 Prerequisites', level=2)
    document.add_paragraph('• Python 3.9 or higher', style='List Bullet')
    document.add_paragraph('• Internet connection (for LLM API calls)', style='List Bullet')
    document.add_paragraph('• Google Gemini API Key (optional, for AI reasoning)', style='List Bullet')

    document.add_heading('3.2 Installation Steps', level=2)
    
    step1 = document.add_paragraph("Step 1: Install Required Libraries", style='Heading 3')
    code1 = document.add_paragraph('pip install streamlit pandas google-generativeai', style='Quote')
    
    step2 = document.add_paragraph("Step 2: Navigate to Project Directory", style='Heading 3')
    code2 = document.add_paragraph('cd fraud_poc', style='Quote')
    
    step3 = document.add_paragraph("Step 3: Launch the Application", style='Heading 3')
    code3 = document.add_paragraph('python -m streamlit run ui_app.py', style='Quote')
    
    step4 = document.add_paragraph("Step 4: Access the Dashboard", style='Heading 3')
    document.add_paragraph("Open your browser and navigate to: http://localhost:8501")

    # --- 4. User Interface Guide ---
    document.add_page_break()
    document.add_heading('4. User Interface Guide', level=1)
    
    document.add_heading('4.1 Sidebar Controls', level=2)
    document.add_paragraph(
        "The left sidebar contains system configuration and status information:"
    )
    document.add_paragraph('• Gemini API Key Input: Optional field for enabling AI reasoning', style='List Bullet')
    document.add_paragraph('• System Status: Shows agent online status', style='List Bullet')
    document.add_paragraph('• Rules Loaded: Count of active fraud detection rules', style='List Bullet')
    document.add_paragraph('• Policies Monitored: Total number of policies in the system', style='List Bullet')

    document.add_heading('4.2 Main Dashboard', level=2)
    document.add_paragraph('The main panel provides:')
    document.add_paragraph('• Claim Selection Dropdown: Choose a claim ID to analyze', style='List Bullet')
    document.add_paragraph('• Analyze Button: Triggers the fraud detection process', style='List Bullet')
    document.add_paragraph('• Results Panel: Displays risk assessment after analysis', style='List Bullet')
    document.add_paragraph('• Recent Claims Queue: Table showing all pending claims', style='List Bullet')

    document.add_heading('4.3 Results Interpretation', level=2)
    document.add_paragraph("After running an analysis, you will see:")
    
    document.add_paragraph()
    p = document.add_paragraph()
    p.add_run("Risk Level Badge: ").bold = True
    p.add_run("Color-coded indicator (Green=Low, Yellow=Medium, Orange=High, Red=Critical)")
    
    p = document.add_paragraph()
    p.add_run("Fraud Probability: ").bold = True
    p.add_run("Percentage score indicating likelihood of fraudulent activity")
    
    p = document.add_paragraph()
    p.add_run("Recommended Action: ").bold = True
    p.add_run("Next steps (Auto-Process, Review, Manual Investigation, SIU Referral)")
    
    p = document.add_paragraph()
    p.add_run("AI Reasoning: ").bold = True
    p.add_run("Natural language explanation of why the claim was flagged")
    
    p = document.add_paragraph()
    p.add_run("Risk Factors: ").bold = True
    p.add_run("List of specific rule violations detected")
    
    p = document.add_paragraph()
    p.add_run("Workflow Diagram: ").bold = True
    p.add_run("Visual flowchart showing the decision path")

    # --- 5. Agent Capabilities ---
    document.add_page_break()
    document.add_heading('5. Agent Capabilities', level=1)
    
    document.add_heading('5.1 Rule-Based Detection', level=2)
    document.add_paragraph("The agent evaluates claims against configurable rules:")
    
    rules_table = document.add_table(rows=6, cols=3)
    rules_table.style = 'Light Grid Accent 1'
    
    headers = rules_table.rows[0].cells
    headers[0].text = 'Rule'
    headers[1].text = 'Condition'
    headers[2].text = 'Risk Score'
    
    rules_data = [
        ('Contestable Period', 'Claim within 2 years of policy start', '40'),
        ('High Value', 'Claim amount > $1,000,000', '30'),
        ('Suspicious Cause', 'Vague or high-risk cause of death', '50'),
        ('Foreign Death', 'Death occurred overseas', '60'),
        ('Lapsed Policy', 'Policy premium not current', '100')
    ]
    
    for idx, (rule, condition, score) in enumerate(rules_data, 1):
        cells = rules_table.rows[idx].cells
        cells[0].text = rule
        cells[1].text = condition
        cells[2].text = score

    document.add_heading('5.2 Behavioral Analysis', level=2)
    document.add_paragraph("Beyond static rules, the agent performs:")
    document.add_paragraph('• Cross-referencing medical history with claim timing', style='List Bullet')
    document.add_paragraph('• Deviation detection (claim amount vs. policy coverage)', style='List Bullet')
    document.add_paragraph('• Policy status validation', style='List Bullet')

    # --- 6. AI/LLM Integration ---
    document.add_page_break()
    document.add_heading('6. AI/LLM Integration', level=1)
    
    document.add_heading('6.1 How It Works', level=2)
    document.add_paragraph(
        "The system features a hybrid reasoning engine:"
    )
    
    p = document.add_paragraph()
    p.add_run("With Gemini API Key: ").bold = True
    p.add_run(
        "The agent sends claim context and risk assessment to Google's Gemini 2.0 Flash-Lite model. "
        "The LLM generates a unique, contextual explanation that reads like a human investigator's notes."
    )
    
    p = document.add_paragraph()
    p.add_run("Without API Key: ").bold = True
    p.add_run(
        "The system uses an advanced template engine that produces professional-sounding reports "
        "by dynamically assembling pre-written phrases based on detected violations."
    )

    document.add_heading('6.2 Obtaining a Gemini API Key', level=2)
    document.add_paragraph("To enable AI reasoning:")
    document.add_paragraph('1. Visit Google AI Studio (https://aistudio.google.com)', style='List Number')
    document.add_paragraph('2. Sign in with your Google account', style='List Number')
    document.add_paragraph('3. Navigate to "Get API Key"', style='List Number')
    document.add_paragraph('4. Create a new API key for your project', style='List Number')
    document.add_paragraph('5. Copy the key and paste it into the sidebar input field', style='List Number')

    document.add_heading('6.3 Privacy & Security', level=2)
    document.add_paragraph(
        "Important: When using the Gemini API, claim data is sent to Google's servers for processing. "
        "Ensure compliance with your organization's data privacy policies before enabling this feature. "
        "API keys are handled securely via Streamlit's password input and are not stored in code or logs."
    )

    document.add_heading('6.4 Understanding the LLM Prompt (Technical Deep Dive)', level=2)
    
    p = document.add_paragraph()
    p.add_run("What Gets Sent to Gemini?").bold = True
    
    document.add_paragraph(
        "The system sends a structured text prompt to the Gemini API that includes:"
    )
    document.add_paragraph('1. Role Definition: Tells the LLM to act as an expert fraud analyst', style='List Number')
    document.add_paragraph('2. Task Description: Explains what analysis is needed', style='List Number')
    document.add_paragraph('3. Claim Context: All relevant claim and policy details', style='List Number')
    document.add_paragraph('4. Risk Assessment Results: What the rule engine already calculated', style='List Number')
    document.add_paragraph('5. Output Instructions: How to format the response', style='List Number')
    
    document.add_paragraph()
    p = document.add_paragraph()
    p.add_run("Example Prompt for Claim C1002:").bold = True
    
    example_prompt = """
You are an expert SIU (Special Investigations Unit) Fraud Analyst.
Analyze the following insurance claim and explain the fraud risk assessment.

CLAIM CONTEXT:
- Claim ID: C1002
- Cause of Death: Accidental Fall
- Location: Remote Hiking Trail
- Claim Amount: $2,000,000
- Policy Start Date: 2023-09-01
- Policy Status: Active

RISK ASSESSMENT:
- Calculated Risk Level: Critical
- Fraud Probability: 90%
- Detected Violations: ['High Value Claim', 'Contestable Period', 'Non-Disclosure Risk']

INSTRUCTIONS:
Write a concise, professional paragraph (max 3 sentences) explaining WHY this claim is flagged.
Focus on the correlation between the policy age, amount, and cause of death.
Do not mention "AI" or "Model". Write as a human investigator.
    """
    document.add_paragraph(example_prompt, style='Quote')
    
    document.add_paragraph()
    p = document.add_paragraph()
    p.add_run("What Gemini Returns:").bold = True
    
    example_response = """
"This claim warrants immediate investigation due to the substantial $2M payout request 
occurring within the contestable period. The accidental fall at a remote hiking trail, 
combined with the policy's recent inception date of September 2023, presents multiple 
red flags that require thorough verification of circumstances and beneficiary relationships. 
The convergence of high claim value and limited policy history suggests potential 
pre-existing intent."
    """
    document.add_paragraph(example_response, style='Quote')
    
    document.add_paragraph()
    p = document.add_paragraph()
    p.add_run("Key Points:").bold = True
    
    document.add_paragraph('• The LLM does NOT make the fraud decision - it only explains the decision already made by the rule engine', style='List Bullet')
    document.add_paragraph('• All risk scoring is done by Python code - the LLM just makes it sound human', style='List Bullet')
    document.add_paragraph('• The prompt uses f-strings to dynamically insert real data from JSON files', style='List Bullet')
    document.add_paragraph('• This technique is called "Prompt Engineering" - crafting input text to get optimal output', style='List Bullet')

    # --- 7. Technical Architecture ---
    document.add_page_break()
    document.add_heading('7. Technical Architecture', level=1)
    
    document.add_heading('7.1 Core Logic Flow', level=2)
    code_flow = """
1. User selects Claim ID
2. Agent fetches claim data from claims.json
3. Agent fetches policy data from policies.json
4. Agent loads fraud rules from insurance_rules.json
5. Rule evaluation engine processes each rule
6. Behavioral heuristics are applied
7. Risk score is calculated
8. Risk level is determined (Low/Medium/High/Critical)
9. LLM generates reasoning (if API key present)
10. Results are displayed in UI with workflow diagram
    """
    document.add_paragraph(code_flow, style='Quote')

    document.add_heading('7.2 Key Functions', level=2)
    
    func_table = document.add_table(rows=5, cols=2)
    func_table.style = 'Light Grid Accent 1'
    
    headers = func_table.rows[0].cells
    headers[0].text = 'Function'
    headers[1].text = 'Purpose'
    
    funcs = [
        ('evaluate_rules()', 'Iterates through rules and checks conditions'),
        ('calculate_days_since_inception()', 'Determines policy age for contestability'),
        ('generate_llm_analysis()', 'Calls Gemini API or uses template'),
        ('run_analysis()', 'Main orchestration function')
    ]
    
    for idx, (func, purpose) in enumerate(funcs, 1):
        cells = func_table.rows[idx].cells
        cells[0].text = func
        cells[1].text = purpose

    # --- 8. Sample Workflows ---
    document.add_page_break()
    document.add_heading('8. Sample Workflows', level=1)
    
    document.add_heading('Scenario 1: Low Risk Claim', level=2)
    document.add_paragraph("Claim C1001: $500,000 claim on 8-year-old policy, natural causes, hospital death")
    document.add_paragraph("Expected Result:")
    document.add_paragraph('• Risk Level: Low', style='List Bullet')
    document.add_paragraph('• Probability: ~5%', style='List Bullet')
    document.add_paragraph('• Action: Auto-Process', style='List Bullet')
    document.add_paragraph('• Reasoning: "Claim aligns with standard policy provisions..."', style='List Bullet')

    document.add_heading('Scenario 2: Critical Risk Claim', level=2)
    document.add_paragraph("Claim C1002: $2M claim on 2-month-old policy, accidental fall, remote location")
    document.add_paragraph("Expected Result:")
    document.add_paragraph('• Risk Level: Critical', style='List Bullet')
    document.add_paragraph('• Probability: ~85%', style='List Bullet')
    document.add_paragraph('• Action: SIU Referral', style='List Bullet')
    document.add_paragraph('• Risk Factors: Contestable Period, High Value, Non-Disclosure Risk', style='List Bullet')

    # --- 9. Troubleshooting ---
    document.add_page_break()
    document.add_heading('9. Troubleshooting', level=1)
    
    troubleshoot_table = document.add_table(rows=4, cols=2)
    troubleshoot_table.style = 'Light Grid Accent 1'
    
    headers = troubleshoot_table.rows[0].cells
    headers[0].text = 'Issue'
    headers[1].text = 'Solution'
    
    issues = [
        ('Streamlit not found', 'Run: pip install streamlit'),
        ('LLM Error displayed', 'Check API key validity and internet connection'),
        ('No claims showing', 'Verify claims.json exists in /data folder')
    ]
    
    for idx, (issue, solution) in enumerate(issues, 1):
        cells = troubleshoot_table.rows[idx].cells
        cells[0].text = issue
        cells[1].text = solution

    # --- 10. Future Enhancements ---
    document.add_page_break()
    document.add_heading('10. Future Enhancements', level=1)
    
    document.add_heading('Phase 2: Machine Learning', level=2)
    document.add_paragraph('• Isolation Forest for anomaly detection', style='List Bullet')
    document.add_paragraph('• Vector embeddings for semantic similarity', style='List Bullet')
    
    document.add_heading('Phase 3: Multi-Agent System', level=2)
    document.add_paragraph('• Risk Scoring Agent', style='List Bullet')
    document.add_paragraph('• OSINT Investigation Agent', style='List Bullet')
    document.add_paragraph('• Decision Synthesis Agent', style='List Bullet')
    
    document.add_heading('Phase 4: Production Features', level=2)
    document.add_paragraph('• FastAPI REST endpoints', style='List Bullet')
    document.add_paragraph('• Medical record NLP parsing', style='List Bullet')
    document.add_paragraph('• Hospital API integration', style='List Bullet')
    document.add_paragraph('• Complete audit trail logging', style='List Bullet')

    # --- Footer ---
    section = document.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.text = "Life Insurance Fraud Detection System v2.0 | Confidential - Internal Use Only"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save
    file_path = os.path.join(os.path.dirname(__file__), 'Life_Insurance_Fraud_Complete_Guide.docx')
    document.save(file_path)
    print(f"Comprehensive guide saved to: {file_path}")

if __name__ == "__main__":
    create_comprehensive_user_guide()
